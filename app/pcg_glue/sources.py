"""Source resolver — text / URL / any-file -> (claim, evidence list).

Supports: .pdf, .txt, .md, .docx, .csv, .xlsx, .xls, .json, .html, .htm

A single resolve_file() entrypoint auto-detects format by extension. Each
parser produces clean text chunks; chunking is uniform (1200 char windows
with 100 char overlap). Evidence carries title/publisher/url metadata so
the Proposer can cite the right chunk.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional
import json
import re


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class Evidence:
    text: str
    title: str = ""
    url: str = ""
    publisher: str = ""
    rank: int = 0
    page: Optional[int] = None
    section: str = ""


@dataclass
class ResolvedInput:
    claim: str
    evidence: list[Evidence]
    source_kind: str
    source_label: str
    raw_text_preview: str = ""


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def _chunk(text: str, max_chars: int = 1200, overlap: int = 100) -> list[str]:
    """Sliding-window chunker. Collapses whitespace, drops empty chunks."""
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]
    out = []
    i = 0
    step = max(1, max_chars - overlap)
    while i < len(text):
        out.append(text[i:i + max_chars])
        i += step
    return out


# ---------------------------------------------------------------------------
# Free-text resolver
# ---------------------------------------------------------------------------

def resolve_text(question: str, context: str = "") -> ResolvedInput:
    if not question.strip():
        raise ValueError("Question is empty.")
    if context.strip():
        chunks = _chunk(context)
        evs = [
            Evidence(text=c, title=f"Pasted context (chunk {i+1})",
                     publisher="user_input", rank=i)
            for i, c in enumerate(chunks)
        ]
    else:
        evs = [Evidence(text="(no evidence supplied)", title="No evidence",
                        publisher="user_input", rank=0)]
    return ResolvedInput(
        claim=question.strip(),
        evidence=evs,
        source_kind="text",
        source_label="Free-form text" + (" + pasted context" if context.strip() else ""),
        raw_text_preview=(context[:400] + "…") if len(context) > 400 else context,
    )


# ---------------------------------------------------------------------------
# URL resolver
# ---------------------------------------------------------------------------

def resolve_url(question: str, url: str, *, timeout: float = 15.0) -> ResolvedInput:
    if not question.strip():
        raise ValueError("Question is empty.")
    if not url.startswith(("http://", "https://")):
        raise ValueError("URL must start with http:// or https://")

    import httpx
    from bs4 import BeautifulSoup

    headers = {"User-Agent": "Mozilla/5.0 (PCG-MAS demo)"}
    with httpx.Client(timeout=timeout, follow_redirects=True, headers=headers) as client:
        resp = client.get(url)
        resp.raise_for_status()
        html = resp.text

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
        tag.decompose()
    title = (soup.title.string or "").strip() if soup.title else url
    raw = " ".join(soup.stripped_strings)
    chunks = _chunk(raw)

    evs = [
        Evidence(text=c, title=title or url, url=url, publisher="web", rank=i)
        for i, c in enumerate(chunks)
    ] or [Evidence(text="(no extractable text)", title=title or url,
                   url=url, publisher="web", rank=0)]
    return ResolvedInput(
        claim=question.strip(),
        evidence=evs,
        source_kind="url",
        source_label=f"{title or url} ({len(evs)} chunks)",
        raw_text_preview=raw[:400] + ("…" if len(raw) > 400 else ""),
    )


# ---------------------------------------------------------------------------
# File parsers — one per supported extension
# ---------------------------------------------------------------------------

def _parse_pdf(p: Path) -> list[tuple[Optional[int], str, str]]:
    """Return list of (page_num, section, text)."""
    from pypdf import PdfReader
    reader = PdfReader(str(p))
    out = []
    for i, page in enumerate(reader.pages):
        try:
            out.append((i + 1, f"page {i + 1}", page.extract_text() or ""))
        except Exception:
            out.append((i + 1, f"page {i + 1}", ""))
    return out


def _parse_txt(p: Path) -> list[tuple[Optional[int], str, str]]:
    txt = p.read_text(encoding="utf-8", errors="replace")
    return [(None, "", txt)]


def _parse_md(p: Path) -> list[tuple[Optional[int], str, str]]:
    """Split markdown by H1/H2 headings if present; otherwise treat as one blob."""
    txt = p.read_text(encoding="utf-8", errors="replace")
    sections = re.split(r"\n(?=#{1,2}\s+)", txt)
    if len(sections) <= 1:
        return [(None, "", txt)]
    out = []
    for sec in sections:
        m = re.match(r"#{1,2}\s+(.+?)\n", sec)
        title = m.group(1).strip() if m else ""
        out.append((None, title, sec))
    return out


def _parse_docx(p: Path) -> list[tuple[Optional[int], str, str]]:
    from docx import Document  # python-docx
    doc = Document(str(p))
    blocks = []
    current_section = ""
    buffer = []
    for para in doc.paragraphs:
        text = para.text or ""
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style and text.strip():
            if buffer:
                blocks.append((None, current_section, "\n".join(buffer)))
                buffer = []
            current_section = text.strip()
        elif text.strip():
            buffer.append(text)
    if buffer:
        blocks.append((None, current_section, "\n".join(buffer)))
    # Tables: dump cell contents row-major
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            blocks.append((None, f"table {ti + 1}", "\n".join(rows)))
    return blocks or [(None, "", "")]


def _parse_csv(p: Path) -> list[tuple[Optional[int], str, str]]:
    """Render CSV as a series of compact text records (one per row)."""
    import csv
    rows = []
    with p.open(newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        try:
            header = next(reader)
        except StopIteration:
            return [(None, "", "")]
        for ri, row in enumerate(reader, start=1):
            kv = " | ".join(
                f"{h}={v}" for h, v in zip(header, row) if v.strip()
            )
            if kv:
                rows.append(f"[row {ri}] {kv}")
    blob = "\n".join(rows)
    return [(None, "csv data", blob)]


def _parse_xlsx(p: Path) -> list[tuple[Optional[int], str, str]]:
    """Read every sheet via openpyxl; render each as compact rows."""
    from openpyxl import load_workbook
    wb = load_workbook(str(p), data_only=True, read_only=True)
    out = []
    for sname in wb.sheetnames:
        sheet = wb[sname]
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        header = [str(c) if c is not None else "" for c in rows[0]]
        body = []
        for ri, row in enumerate(rows[1:], start=1):
            kv = " | ".join(
                f"{h}={v}" for h, v in zip(header, row)
                if v is not None and str(v).strip()
            )
            if kv:
                body.append(f"[row {ri}] {kv}")
        if body:
            out.append((None, f"sheet {sname}", "\n".join(body)))
    return out or [(None, "", "")]


def _parse_json(p: Path) -> list[tuple[Optional[int], str, str]]:
    """Render JSON as pretty-printed text; one chunk per top-level key for dicts."""
    obj = json.loads(p.read_text(encoding="utf-8", errors="replace"))
    if isinstance(obj, dict):
        out = []
        for k, v in obj.items():
            out.append((None, str(k), json.dumps(v, indent=2)))
        return out or [(None, "", "")]
    return [(None, "", json.dumps(obj, indent=2))]


def _parse_html(p: Path) -> list[tuple[Optional[int], str, str]]:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(p.read_text(encoding="utf-8", errors="replace"), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return [(None, soup.title.string.strip() if soup.title else "",
             " ".join(soup.stripped_strings))]


# Registry — single source of truth for which extensions we support
_PARSER_REGISTRY: dict[str, callable] = {
    ".pdf":  _parse_pdf,
    ".txt":  _parse_txt,
    ".md":   _parse_md,
    ".markdown": _parse_md,
    ".docx": _parse_docx,
    ".csv":  _parse_csv,
    ".xlsx": _parse_xlsx,
    ".xls":  _parse_xlsx,
    ".json": _parse_json,
    ".html": _parse_html,
    ".htm":  _parse_html,
}

SUPPORTED_EXTENSIONS: list[str] = sorted(_PARSER_REGISTRY.keys())


def resolve_file(question: str, file_path: str | Path) -> ResolvedInput:
    """Auto-detect file type by extension, parse, chunk, build evidence list."""
    if not question.strip():
        raise ValueError("Question is empty.")
    p = Path(file_path)
    if not p.exists():
        raise ValueError(f"File not found: {p}")
    ext = p.suffix.lower()
    parser = _PARSER_REGISTRY.get(ext)
    if parser is None:
        raise ValueError(
            f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    try:
        blocks = parser(p)
    except Exception as e:
        raise ValueError(f"Could not parse {p.name} ({ext}): {type(e).__name__}: {e}")

    evs: list[Evidence] = []
    total_chars = 0
    for page_num, section, text in blocks:
        for chunk in _chunk(text):
            if chunk.strip():
                evs.append(Evidence(
                    text=chunk,
                    title=f"{p.name}" + (f" · {section}" if section else ""),
                    publisher=f"upload:{ext.lstrip('.')}",
                    rank=len(evs),
                    page=page_num,
                    section=section,
                ))
                total_chars += len(chunk)

    if not evs:
        evs = [Evidence(
            text=f"(no extractable text from {p.name})",
            title=p.name, publisher=f"upload:{ext.lstrip('.')}", rank=0,
        )]

    preview_text = " ".join(ev.text for ev in evs[:2])[:400]
    return ResolvedInput(
        claim=question.strip(),
        evidence=evs,
        source_kind=f"file/{ext.lstrip('.')}",
        source_label=f"{p.name} · {len(evs)} chunks · ~{total_chars // 1000}k chars",
        raw_text_preview=preview_text + ("…" if total_chars > 400 else ""),
    )


# Backwards-compat alias for the old PDF-only entrypoint
def resolve_pdf(question: str, pdf_path: str | Path) -> ResolvedInput:
    return resolve_file(question, pdf_path)


# ---------------------------------------------------------------------------
# Retrieval helper
# ---------------------------------------------------------------------------

def top_k(resolved: ResolvedInput, k: int) -> ResolvedInput:
    """Trim evidence to top_k by simple keyword overlap (BM25-lite)."""
    if k >= len(resolved.evidence):
        return resolved
    q_tokens = set(re.findall(r"\w+", resolved.claim.lower()))
    if not q_tokens:
        return resolved
    scored = []
    for ev in resolved.evidence:
        ev_tokens = set(re.findall(r"\w+", ev.text.lower()))
        scored.append((len(q_tokens & ev_tokens), ev))
    scored.sort(key=lambda x: -x[0])
    chosen = [ev for _, ev in scored[:k]]
    for i, ev in enumerate(chosen):
        ev.rank = i
    return ResolvedInput(
        claim=resolved.claim,
        evidence=chosen,
        source_kind=resolved.source_kind,
        source_label=resolved.source_label + f" · top-{k} retrieved",
        raw_text_preview=resolved.raw_text_preview,
    )
